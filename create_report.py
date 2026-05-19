import os
from docx import Document
from docx.shared import Pt, Inches

def create_word():
    doc = Document()
    
    title = doc.add_heading('BÁO CÁO THỰC TẬP/DỰ ÁN: HỆ THỐNG VNU RESEARCH API', 0)
    title.alignment = 1 # Center
    
    doc.add_heading('1. Giới thiệu', level=1)
    
    doc.add_heading('1.1. Phát biểu bài toán', level=2)
    doc.add_paragraph('Công việc được giao là xây dựng và tối ưu hoá hệ thống API tìm kiếm và quản lý đề tài nghiên cứu khoa học của trường đại học (VNU Research Web API). Hệ thống cần đảm bảo khả năng tra cứu nhanh chóng, chính xác các đề tài dựa trên từ khóa, đối tượng (Giảng viên, Sinh viên) và khả năng mở rộng trong tương lai. Bài toán đặt ra yêu cầu xây dựng một kiến trúc Backend hiệu năng cao, chuẩn bị sẵn sàng cho việc đóng gói và triển khai (Containerization).')
    
    doc.add_heading('1.2. Công nghệ và công cụ sử dụng', level=2)
    p = doc.add_paragraph()
    p.add_run('1. FastAPI: ').bold = True
    p.add_run('Framework Python dùng để xây dựng API.\n')
    p.add_run('- Ưu điểm: Hiệu năng cực cao nhờ cơ chế ASGI, hỗ trợ lập trình bất đồng bộ (async/await), tự động sinh tài liệu Swagger UI/ReDoc giúp quá trình kiểm thử và giao tiếp API diễn ra thuận lợi.\n')
    p.add_run('- Lý do lựa chọn: Tốc độ xử lý nhanh và phù hợp với mô hình microservices hoặc kiến trúc SOA tinh gọn.')
    
    p = doc.add_paragraph()
    p.add_run('2. PostgreSQL: ').bold = True
    p.add_run('Hệ quản trị cơ sở dữ liệu quan hệ.\n')
    p.add_run('- Ưu điểm: Ổn định, mạnh mẽ, và đặc biệt hỗ trợ công cụ Full Text Search tích hợp sẵn thông qua kiểu dữ liệu TSVECTOR và extension pg_trgm.\n')
    p.add_run('- Lý do lựa chọn: Cho phép biến database thành một Search Engine cơ bản hiệu quả, không cần phụ thuộc vào ElasticSearch ở giai đoạn đầu, tiết kiệm tài nguyên vận hành.')

    p = doc.add_paragraph()
    p.add_run('3. Redis: ').bold = True
    p.add_run('Hệ thống lưu trữ dữ liệu in-memory.\n')
    p.add_run('- Ưu điểm: Đọc/ghi siêu nhanh.\n')
    p.add_run('- Lý do lựa chọn: Được sử dụng để phân tán Cache kết quả tìm kiếm nhằm giảm tải cho CSDL chính, đồng thời quản lý trạng thái bảo mật (ví dụ: JWT Blacklist).')

    p = doc.add_paragraph()
    p.add_run('4. Docker & Docker Compose: ').bold = True
    p.add_run('Công cụ Container hóa.\n')
    p.add_run('- Ưu điểm: Đóng gói môi trường đồng nhất, triệt tiêu lỗi "works on my machine".\n')
    p.add_run('- Lý do lựa chọn: Dễ dàng cấu hình và liên kết các dịch vụ (API, PostgreSQL, Redis) thông qua docker-compose, chuẩn bị cho CI/CD và triển khai thực tế.')

    doc.add_heading('2. Công việc triển khai', level=1)
    doc.add_paragraph('Trong quá trình tham gia phát triển hệ thống, các hạng mục công việc đã được triển khai bao gồm:')
    
    doc.add_heading('2.1. Thiết kế và Xây dựng Cơ sở dữ liệu', level=2)
    doc.add_paragraph('- Xây dựng cấu trúc các bảng dữ liệu (schema) trên PostgreSQL.')
    doc.add_paragraph('- Áp dụng kỹ thuật Full Text Search (FTS): Tạo các generated columns lưu trữ TSVECTOR và cấu hình GIN index, B-Tree index cho các trường phân loại (năm, đối tượng mục tiêu).')
    
    doc.add_heading('2.2. Xây dựng API Layer (Backend)', level=2)
    doc.add_paragraph('- Xây dựng các endpoints RESTful bằng FastAPI.')
    doc.add_paragraph('- Tối ưu hoá luồng tìm kiếm (Search Flow): Thay thế thuật toán tìm kiếm 2 bước cũ thành thuật toán single-pass, kết hợp FTS và tính điểm số tương đồng (Trigram similarity) nhằm đem lại kết quả chính xác nhất.')
    doc.add_paragraph('- Tích hợp Redis Cache: Xây dựng cơ chế chuẩn hoá cache keys để đảm bảo dữ liệu luôn được "fresh" và giảm tỷ lệ truy vấn lặp xuống Database.')
    doc.add_paragraph('- Xây dựng hệ thống bảo mật: Xác thực dựa trên JWT (Access Token, Refresh Token) và áp dụng cơ chế Rate Limiting để chống Spam/DDoS.')

    doc.add_heading('2.3. Đóng gói và DevOps', level=2)
    doc.add_paragraph('- Container hóa dự án bằng Docker (viết Dockerfile tối ưu kích thước cho Python backend).')
    doc.add_paragraph('- Thiết lập môi trường Docker Compose gồm đầy đủ các services: api, postgres, redis.')
    doc.add_paragraph('- Xây dựng entrypoint script tự động hoá việc kiểm tra kết nối CSDL, chạy database migrations và cấp dữ liệu mẫu (seed data) khi khởi chạy container.')
    doc.add_paragraph('- Cấu hình hệ thống giám sát (Observability Layer) sử dụng Prometheus và Grafana để thu thập metrics (nhưng không làm nghẽn luồng xử lý của API).')

    doc.add_heading('3. Các kết quả đã đạt được', level=1)
    
    doc.add_heading('3.1. Sản phẩm thu được', level=2)
    doc.add_paragraph('- Một hệ thống Web API tìm kiếm hoàn chỉnh, có hiệu năng cao, ổn định và sẵn sàng triển khai trên môi trường Production.')
    doc.add_paragraph('- Giao diện sử dụng (dành cho phía nhà phát triển / tích hợp): Hệ thống cung cấp trang tài liệu API tự động qua Swagger UI (/docs) hiển thị rõ các endpoints như `/api/v1/search`, `/api/v1/projects` cùng cấu trúc tham số (request params, body) và các mô hình phản hồi (response schemas).')
    
    doc.add_heading('3.2. Bản phân tích thiết kế hệ thống', level=2)
    doc.add_paragraph('Hệ thống được thiết kế theo kiến trúc Service-Oriented Architecture (SOA), bao gồm các phân lớp cụ thể:')
    doc.add_paragraph('- API Layer (FastAPI): Gateway xử lý HTTP requests, định tuyến, phân quyền và giới hạn lưu lượng.')
    doc.add_paragraph('- Data Layer (PostgreSQL): Đóng vai trò vừa là bộ nhớ lưu trữ bền vững (Relational Database) vừa là Search Engine (thông qua FTS).')
    doc.add_paragraph('- Cache Layer (Redis): Tăng tốc độ phản hồi và bảo vệ DB bằng Distributed Cache, quản lý phiên qua Token Blacklist.')
    doc.add_paragraph('- Giới hạn và Mở rộng: Hiện tại thiết kế Search Engine nằm trên PostgreSQL phù hợp cho lưu lượng ban đầu. Nếu vượt mốc giới hạn, kiến trúc hỗ trợ mở rộng bằng cách tách luồng tìm kiếm sang Elasticsearch. Khả năng scale-out theo chiều ngang của API Layer là không giới hạn vì kiến trúc FastAPI được thiết kế dạng stateless.')

    doc.add_heading('3.3. Kết quả Test và Khắc phục lỗi', level=2)
    doc.add_paragraph('- Kết quả Test API: Đảm bảo bộ lọc hoạt động chính xác qua các tham số query param.')
    doc.add_paragraph('- Phân tích lỗi đã khắc phục:')
    doc.add_paragraph('  + Lỗi sai lệch bộ lọc tìm kiếm: Khi truyền bộ lọc "Sinh viên" hoặc "Giảng viên", API trả về toàn bộ kết quả. Nguyên nhân do xung đột Cache Key (chưa chuẩn hoá param) và logic query trong backend bị sai lệch tham chiếu. Khắc phục: Chuẩn hoá lại logic tạo custom key cho Redis và viết lại truy vấn DB.')
    doc.add_paragraph('  + Lỗi khởi động Docker: Xung đột volume mount và DB chưa sẵn sàng nhận kết nối từ API. Khắc phục: Sử dụng wait-for-it pattern trong bash script (entrypoint.sh) để đảm bảo trình tự boot services.')

    doc.add_heading('4. Các phụ lục', level=1)
    
    doc.add_heading('4.1. Phân chia công việc', level=2)
    doc.add_paragraph('- Do đây là báo cáo cá nhân, sinh viên đảm nhận toàn bộ các vai trò từ thiết kế kiến trúc, triển khai luồng Backend API, cấu hình Database cho tới việc thiết lập môi trường Docker.')
    
    doc.add_heading('4.2. Tài liệu tham khảo', level=2)
    doc.add_paragraph('1. Tài liệu chính thức của FastAPI: https://fastapi.tiangolo.com/')
    doc.add_paragraph('2. PostgreSQL Full Text Search & pg_trgm documentation.')
    doc.add_paragraph('3. Kiến trúc dự án: Dựa trên đặc tả nội bộ (ARCHITECTURE.md).')

    doc.add_heading('4.3. Ký hiệu viết tắt', level=2)
    doc.add_paragraph('- API: Application Programming Interface')
    doc.add_paragraph('- FTS: Full Text Search')
    doc.add_paragraph('- JWT: JSON Web Token')
    doc.add_paragraph('- SOA: Service-Oriented Architecture')
    
    doc.save('Bao_cao_he_thong.docx')

def create_html():
    html_content = """<!DOCTYPE html>
<html lang="vi">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Báo Cáo Thiết Kế Hệ Thống</title>
    <style>
        body { font-family: Arial, sans-serif; line-height: 1.6; margin: 40px auto; max-width: 800px; padding: 0 20px; color: #333; }
        h1 { text-align: center; color: #2c3e50; }
        h2 { color: #2980b9; border-bottom: 2px solid #eee; padding-bottom: 5px; margin-top: 30px; }
        h3 { color: #34495e; }
        p { margin-bottom: 15px; }
        ul { margin-bottom: 15px; }
        .bold { font-weight: bold; }
    </style>
</head>
<body>
    <h1>BÁO CÁO THỰC TẬP/DỰ ÁN: HỆ THỐNG VNU RESEARCH API</h1>
    
    <h2>1. Giới thiệu</h2>
    <h3>1.1. Phát biểu bài toán</h3>
    <p>Công việc được giao là xây dựng và tối ưu hoá hệ thống API tìm kiếm và quản lý đề tài nghiên cứu khoa học của trường đại học (VNU Research Web API). Hệ thống cần đảm bảo khả năng tra cứu nhanh chóng, chính xác các đề tài dựa trên từ khóa, đối tượng (Giảng viên, Sinh viên) và khả năng mở rộng trong tương lai. Bài toán đặt ra yêu cầu xây dựng một kiến trúc Backend hiệu năng cao, chuẩn bị sẵn sàng cho việc đóng gói và triển khai (Containerization).</p>
    
    <h3>1.2. Công nghệ và công cụ sử dụng</h3>
    <ul>
        <li><span class="bold">FastAPI:</span> Framework Python dùng để xây dựng API.
            <ul>
                <li>Ưu điểm: Hiệu năng cực cao nhờ cơ chế ASGI, hỗ trợ lập trình bất đồng bộ (async/await), tự động sinh tài liệu Swagger UI/ReDoc giúp quá trình kiểm thử và giao tiếp API diễn ra thuận lợi.</li>
                <li>Lý do lựa chọn: Tốc độ xử lý nhanh và phù hợp với mô hình microservices hoặc kiến trúc SOA tinh gọn.</li>
            </ul>
        </li>
        <li><span class="bold">PostgreSQL:</span> Hệ quản trị cơ sở dữ liệu quan hệ.
            <ul>
                <li>Ưu điểm: Ổn định, mạnh mẽ, và đặc biệt hỗ trợ công cụ Full Text Search tích hợp sẵn thông qua kiểu dữ liệu TSVECTOR và extension pg_trgm.</li>
                <li>Lý do lựa chọn: Cho phép biến database thành một Search Engine cơ bản hiệu quả, không cần phụ thuộc vào ElasticSearch ở giai đoạn đầu, tiết kiệm tài nguyên vận hành.</li>
            </ul>
        </li>
        <li><span class="bold">Redis:</span> Hệ thống lưu trữ dữ liệu in-memory.
            <ul>
                <li>Ưu điểm: Đọc/ghi siêu nhanh.</li>
                <li>Lý do lựa chọn: Được sử dụng để phân tán Cache kết quả tìm kiếm nhằm giảm tải cho CSDL chính, đồng thời quản lý trạng thái bảo mật (ví dụ: JWT Blacklist).</li>
            </ul>
        </li>
        <li><span class="bold">Docker & Docker Compose:</span> Công cụ Container hóa.
            <ul>
                <li>Ưu điểm: Đóng gói môi trường đồng nhất, triệt tiêu lỗi "works on my machine".</li>
                <li>Lý do lựa chọn: Dễ dàng cấu hình và liên kết các dịch vụ (API, PostgreSQL, Redis) thông qua docker-compose, chuẩn bị cho CI/CD và triển khai thực tế.</li>
            </ul>
        </li>
    </ul>

    <h2>2. Công việc triển khai</h2>
    <p>Trong quá trình tham gia phát triển hệ thống, các hạng mục công việc đã được triển khai bao gồm:</p>
    
    <h3>2.1. Thiết kế và Xây dựng Cơ sở dữ liệu</h3>
    <ul>
        <li>Xây dựng cấu trúc các bảng dữ liệu (schema) trên PostgreSQL.</li>
        <li>Áp dụng kỹ thuật Full Text Search (FTS): Tạo các generated columns lưu trữ TSVECTOR và cấu hình GIN index, B-Tree index cho các trường phân loại (năm, đối tượng mục tiêu).</li>
    </ul>
    
    <h3>2.2. Xây dựng API Layer (Backend)</h3>
    <ul>
        <li>Xây dựng các endpoints RESTful bằng FastAPI.</li>
        <li>Tối ưu hoá luồng tìm kiếm (Search Flow): Thay thế thuật toán tìm kiếm 2 bước cũ thành thuật toán single-pass, kết hợp FTS và tính điểm số tương đồng (Trigram similarity) nhằm đem lại kết quả chính xác nhất.</li>
        <li>Tích hợp Redis Cache: Xây dựng cơ chế chuẩn hoá cache keys để đảm bảo dữ liệu luôn được "fresh" và giảm tỷ lệ truy vấn lặp xuống Database.</li>
        <li>Xây dựng hệ thống bảo mật: Xác thực dựa trên JWT (Access Token, Refresh Token) và áp dụng cơ chế Rate Limiting để chống Spam/DDoS.</li>
    </ul>

    <h3>2.3. Đóng gói và DevOps</h3>
    <ul>
        <li>Container hóa dự án bằng Docker (viết Dockerfile tối ưu kích thước cho Python backend).</li>
        <li>Thiết lập môi trường Docker Compose gồm đầy đủ các services: api, postgres, redis.</li>
        <li>Xây dựng entrypoint script tự động hoá việc kiểm tra kết nối CSDL, chạy database migrations và cấp dữ liệu mẫu (seed data) khi khởi chạy container.</li>
        <li>Cấu hình hệ thống giám sát (Observability Layer) sử dụng Prometheus và Grafana để thu thập metrics.</li>
    </ul>

    <h2>3. Các kết quả đã đạt được</h2>
    <h3>3.1. Sản phẩm thu được</h3>
    <ul>
        <li>Một hệ thống Web API tìm kiếm hoàn chỉnh, có hiệu năng cao, ổn định và sẵn sàng triển khai trên môi trường Production.</li>
        <li>Giao diện sử dụng (dành cho phía nhà phát triển / tích hợp): Hệ thống cung cấp trang tài liệu API tự động qua Swagger UI (/docs) hiển thị rõ các endpoints như <code>/api/v1/search</code>, <code>/api/v1/projects</code> cùng cấu trúc tham số và các mô hình phản hồi.</li>
    </ul>
    
    <h3>3.2. Bản phân tích thiết kế hệ thống</h3>
    <p>Hệ thống được thiết kế theo kiến trúc Service-Oriented Architecture (SOA), bao gồm các phân lớp cụ thể:</p>
    <ul>
        <li><span class="bold">API Layer (FastAPI):</span> Gateway xử lý HTTP requests, định tuyến, phân quyền và giới hạn lưu lượng.</li>
        <li><span class="bold">Data Layer (PostgreSQL):</span> Đóng vai trò vừa là bộ nhớ lưu trữ bền vững vừa là Search Engine (thông qua FTS).</li>
        <li><span class="bold">Cache Layer (Redis):</span> Tăng tốc độ phản hồi và bảo vệ DB bằng Distributed Cache, quản lý phiên qua Token Blacklist.</li>
        <li><span class="bold">Giới hạn và Mở rộng:</span> Hiện tại thiết kế Search Engine nằm trên PostgreSQL phù hợp cho lưu lượng ban đầu. Kiến trúc hỗ trợ mở rộng bằng cách tách luồng tìm kiếm sang Elasticsearch trong tương lai.</li>
    </ul>

    <h3>3.3. Kết quả Test và Khắc phục lỗi</h3>
    <ul>
        <li>Kết quả Test API: Đảm bảo bộ lọc hoạt động chính xác qua các tham số query param.</li>
        <li>Phân tích lỗi đã khắc phục:
            <ul>
                <li><span class="bold">Lỗi sai lệch bộ lọc tìm kiếm:</span> Khi truyền bộ lọc "Sinh viên" hoặc "Giảng viên", API trả về toàn bộ kết quả. Nguyên nhân do xung đột Cache Key và logic query trong backend bị sai lệch tham chiếu. Khắc phục: Chuẩn hoá lại logic tạo custom key cho Redis và viết lại truy vấn DB.</li>
                <li><span class="bold">Lỗi khởi động Docker:</span> Xung đột volume mount và DB chưa sẵn sàng nhận kết nối từ API. Khắc phục: Sử dụng wait-for-it pattern trong bash script để đảm bảo trình tự boot services.</li>
            </ul>
        </li>
    </ul>

    <h2>4. Các phụ lục</h2>
    <h3>4.1. Phân chia công việc</h3>
    <p>Sinh viên đảm nhận các vai trò từ thiết kế kiến trúc, triển khai luồng Backend API, cấu hình Database cho tới việc thiết lập môi trường Docker.</p>
    
    <h3>4.2. Tài liệu tham khảo</h3>
    <ul>
        <li>Tài liệu chính thức của FastAPI: https://fastapi.tiangolo.com/</li>
        <li>PostgreSQL Full Text Search & pg_trgm documentation.</li>
        <li>Kiến trúc dự án nội bộ (ARCHITECTURE.md).</li>
    </ul>

    <h3>4.3. Ký hiệu viết tắt</h3>
    <ul>
        <li>API: Application Programming Interface</li>
        <li>FTS: Full Text Search</li>
        <li>JWT: JSON Web Token</li>
        <li>SOA: Service-Oriented Architecture</li>
    </ul>
</body>
</html>"""
    with open('Bao_cao_he_thong.html', 'w', encoding='utf-8') as f:
        f.write(html_content)

if __name__ == "__main__":
    try:
        create_word()
    except ImportError:
        print("Không tìm thấy thư viện python-docx. Đang tạo file .doc dạng văn bản thay thế...")
        # fallback: save html content as .doc extension
        with open('Bao_cao_he_thong.doc', 'w', encoding='utf-8') as f:
            f.write("<html><head><meta charset='utf-8'></head><body><h1>Vui lòng cài python-docx để tạo file .docx chuẩn, đây là bản preview</h1></body></html>")
    create_html()
